import os
import requests
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import CharacterTextSplitter
from langchain.schema import Document
from urllib.parse import urljoin, urlparse
import time
from typing import Set, List

class DocumentIngestor:
    def __init__(self):
        self.text_splitter = CharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=200,
            separator="\n"
        )
        self.visited_urls = set()  
        self.max_pages = 15  
        self.delay = 1  
        
    def load_pdfs(self, pdf_directory="data/pdfs/"):
        
        documents = []
        
        if not os.path.exists(pdf_directory):
            print(f"Directory {pdf_directory} not found")
            return documents
            
        for filename in os.listdir(pdf_directory):
            if filename.endswith('.pdf'):
                pdf_path = os.path.join(pdf_directory, filename)
                try:
                    reader = PdfReader(pdf_path)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text()
                    
                    documents.append(Document(
                        page_content=text,
                        metadata={"source": filename, "type": "pdf"}
                    ))
                    print(f"Loaded PDF: {filename}")
                except Exception as e:
                    print(f"Error loading PDF {filename}: {e}")
        
        return documents
    
    def load_docx_files(self, directory="data/pdfs/"):
        
        documents = []
        
        if not os.path.exists(directory):
            print(f"Directory {directory} not found")
            return documents
            
        for filename in os.listdir(directory):
            if filename.endswith('.docx') and not filename.startswith('~$'):
                docx_path = os.path.join(directory, filename)
                try:
                    doc = DocxDocument(docx_path)
                    
                    text_content = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(paragraph.text.strip())
                    
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                text_content.append(" | ".join(row_text))
                    
                    full_text = "\n".join(text_content)
                    
                    if full_text.strip():
                        documents.append(Document(
                            page_content=full_text,
                            metadata={"source": filename, "type": "docx"}
                        ))
                        print(f"Loaded DOCX: {filename}")
                    else:
                        print(f"Warning: DOCX file {filename} appears to be empty")
                        
                except Exception as e:
                    print(f"Error loading DOCX {filename}: {e}")
        
        return documents
    
    def load_all_documents(self, directory="data/pdfs/"):
        
        documents = []
        
        pdf_docs = self.load_pdfs(directory)
        documents.extend(pdf_docs)
        
        docx_docs = self.load_docx_files(directory)
        documents.extend(docx_docs)
        
        print(f"Total documents loaded: {len(documents)} ({len(pdf_docs)} PDFs, {len(docx_docs)} DOCX)")
        return documents
    
    def _get_internal_links(self, soup, base_url: str) -> List[str]:
        
        links = []
        base_domain = urlparse(base_url).netloc
        
        
        for link in soup.find_all('a', href=True):
            href = link['href']
            
            full_url = urljoin(base_url, href)
            parsed_url = urlparse(full_url)
            
            if parsed_url.netloc == base_domain:
                
                if not any(skip in href.lower() for skip in ['#', 'javascript:', 'mailto:', 'tel:', '.pdf', '.jpg', '.png', '.gif']):
                    links.append(full_url)
        
        return list(set(links))  
    
    def _scrape_single_page(self, url: str) -> Document:
        
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            
            for script in soup(["script", "style", "nav", "header", "footer"]):
                script.decompose()
            
            
            content_selectors = [
                'main', 'article', '.content', '#content', 
                '.main-content', 'section', '.container',
                '.post-content', '.entry-content'
            ]
            
            content_text = ""
            content_found = False
            
            for selector in content_selectors:
                elements = soup.select(selector)
                if elements:
                    for element in elements:
                        text = element.get_text(separator='\n', strip=True)
                        if len(text) > 100:
                            content_text += text + "\n"
                            content_found = True
                    if content_found:
                        break
            

            if not content_found:
                content_text = soup.get_text(separator='\n', strip=True)
            
            
            lines = content_text.split('\n')
            cleaned_lines = []
            for line in lines:
                line = line.strip()
                if len(line) > 10 and not line.startswith(('Home', 'Menu', 'Search', 'Login', 'Contact')):
                    cleaned_lines.append(line)
            
            final_text = '\n'.join(cleaned_lines)
            
            if len(final_text) > 100:
                return Document(
                    page_content=final_text,
                    metadata={"source": url, "type": "webpage", "title": soup.title.string if soup.title else "Unknown"}
                )
            else:
                return None
                
        except Exception as e:
            print(f"Error scraping {url}: {e}")
            return None
    
    def scrape_website_recursive(self, start_url: str = "https://www.angelone.in/support", max_pages: int = 10) -> List[Document]:
        
        documents = []
        self.visited_urls.clear()
        self.max_pages = max_pages
        
        urls_to_visit = [start_url]
        pages_scraped = 0
        
        print(f"Starting recursive scraping from: {start_url}")
        print(f"Maximum pages to scrape: {max_pages}")
        
        while urls_to_visit and pages_scraped < self.max_pages:
            current_url = urls_to_visit.pop(0)
            
            
            if current_url in self.visited_urls:
                continue
                
            print(f"Scraping page {pages_scraped + 1}/{max_pages}: {current_url}")
            self.visited_urls.add(current_url)
            
            
            document = self._scrape_single_page(current_url)
            if document:
                documents.append(document)
                print(f"✅ Successfully scraped: {current_url}")
                
                
                try:
                    response = requests.get(current_url, headers={'User-Agent': 'Mozilla/5.0'}, timeout=10)
                    soup = BeautifulSoup(response.content, 'html.parser')
                    internal_links = self._get_internal_links(soup, current_url)
                    
                    
                    for link in internal_links:
                        if link not in self.visited_urls and link not in urls_to_visit:
                            urls_to_visit.append(link)
                            
                except Exception as e:
                    print(f"Error finding links on {current_url}: {e}")
            else:
                print(f"❌ Failed to scrape: {current_url}")
            
            pages_scraped += 1
            
            
            if pages_scraped < self.max_pages:
                time.sleep(self.delay)
        
        print(f"\n📊 Scraping Summary:")
        print(f"Pages scraped: {pages_scraped}")
        print(f"Documents created: {len(documents)}")
        print(f"URLs visited: {len(self.visited_urls)}")
        
        return documents
    
    def scrape_webpage(self, url: str = "https://www.angelone.in/support") -> List[Document]:
        
        return self.scrape_website_recursive(url, max_pages=1)
    
    def process_documents(self, documents):
        
        if not documents:
            return []
        
        chunks = self.text_splitter.split_documents(documents)
        print(f"Split {len(documents)} documents into {len(chunks)} chunks")
        return chunks

if __name__ == "__main__":
    ingestor = DocumentIngestor()
    
    
    all_file_docs = ingestor.load_all_documents()
    
    
    web_docs = ingestor.scrape_website_recursive(
        start_url="https://www.angelone.in/support", 
        max_pages=10  
    )
    

    all_docs = all_file_docs + web_docs
    
    
    chunks = ingestor.process_documents(all_docs)
    
    print(f"\n🎯 Final Results:")
    print(f"Total documents: {len(all_docs)}")
    print(f"Total chunks: {len(chunks)}")
    
    
    print(f"\n🌐 Scraped URLs:")
    for i, url in enumerate(ingestor.visited_urls, 1):
        print(f"{i}. {url}")